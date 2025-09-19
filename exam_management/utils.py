import io
import docx
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image as PillowImage # Pillow library for image processing

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, Image as ReportLabImage
from reportlab.lib.units import inch

# ==============================================================================
# Constants
# ==============================================================================

THAI_CHOICE_CHARS = [
    'ก', 'ข', 'ค', 'ง', 'จ', 'ฉ', 'ช', 'ซ', 'ฌ', 'ญ', 
    'ฎ', 'ฏ', 'ฐ', 'ฑ', 'ฒ', 'ณ', 'ด', 'ต', 'ถ', 'ท', 
    'ธ', 'น', 'บ', 'ป', 'ผ', 'ฝ', 'พ', 'ฟ', 'ภ', 'ม', 
    'ย', 'ร', 'ล', 'ว', 'ศ', 'ษ', 'ส', 'ห', 'ฬ', 'อ', 'ฮ'
]

# ==============================================================================
# Font Setup for ReportLab (PDF Generation)
# ==============================================================================
try:
    font_path = 'static/fonts/THSarabunNew.ttf' 
    pdfmetrics.registerFont(TTFont('ThaiFont', font_path))
    pdfmetrics.registerFont(TTFont('ThaiFont-Bold', font_path.replace(".ttf", " Bold.ttf")))

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ThaiBody', fontName='ThaiFont', fontSize=12, leading=14))
    styles.add(ParagraphStyle(name='ThaiHeader', fontName='ThaiFont-Bold', fontSize=16, leading=18, spaceAfter=6))
    styles.add(ParagraphStyle(name='ThaiSubHeader', fontName='ThaiFont', fontSize=12, leading=14, spaceAfter=6))
    styles.add(ParagraphStyle(name='ThaiQuestion', fontName='ThaiFont', fontSize=12, leading=14, leftIndent=inch*0.2))
    styles.add(ParagraphStyle(name='ThaiChoice', fontName='ThaiFont', fontSize=12, leading=14, leftIndent=inch*0.4))
    
    FONT_SETUP_SUCCESS = True
except Exception as e:
    print(f"Warning: Could not load Thai font for PDF generation. Error: {e}")
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='ThaiBody', parent=styles['BodyText']))
    styles.add(ParagraphStyle(name='ThaiHeader', parent=styles['h2']))
    styles.add(ParagraphStyle(name='ThaiSubHeader', parent=styles['h3']))
    styles.add(ParagraphStyle(name='ThaiQuestion', parent=styles['BodyText'], leftIndent=inch*0.2))
    styles.add(ParagraphStyle(name='ThaiChoice', parent=styles['BodyText'], leftIndent=inch*0.4))
    FONT_SETUP_SUCCESS = False


# ==============================================================================
# PDF Generation Utility
# ==============================================================================

def generate_pdf_exam(exam, choice_format='thai'):
    """
    Generates a PDF file for a given Exam object, including images.
    """
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # --- Header ---
    p.setFont('ThaiFont-Bold' if FONT_SETUP_SUCCESS else 'Helvetica-Bold', 16)
    p.drawString(inch, height - inch, f"ชุดข้อสอบ: {exam.exam_name}")
    
    p.setFont('ThaiFont' if FONT_SETUP_SUCCESS else 'Helvetica', 12)
    p.drawString(inch, height - inch - 20, f"รายวิชา: {exam.course}")

    p.line(inch, height - inch - 30, width - inch, height - inch - 30)

    # --- Questions ---
    y_position = height - inch - 60
    for i, question in enumerate(exam.questions.all(), 1):
        # Estimate height to check for page break before drawing
        est_height = 50 
        if question.image:
            est_height += 150
        if question.question_type == 'MCQ':
            est_height += len(question.choices.all()) * 20
        
        if y_position - est_height < inch:
            p.showPage()
            y_position = height - inch

        # Draw question text
        question_text = f"{i}. {question.question_text}"
        para = Paragraph(question_text, styles['ThaiQuestion'])
        w, h = para.wrapOn(p, width - 2*inch, height)
        para.drawOn(p, inch, y_position - h)
        y_position -= (h + 10)

        # Draw image if it exists
        if question.image:
            try:
                # Set a max width and let height be proportional
                max_width = 3 * inch
                img = ReportLabImage(question.image.path, width=max_width, height=max_width * 0.75) # Aspect ratio guess
                img.hAlign = 'LEFT'
                
                if y_position - img.drawHeight < inch:
                    p.showPage()
                    y_position = height - inch
                
                img.drawOn(p, inch * 1.2, y_position - img.drawHeight)
                y_position -= (img.drawHeight + 10)
            except Exception as e:
                print(f"Error adding image to PDF: {e}")
                p.drawString(inch * 1.2, y_position - 12, f"[ไม่สามารถแทรกรูปภาพ: {question.image.name}]")
                y_position -= 20

        # Draw choices if MCQ
        if question.question_type == 'MCQ':
            choices = list(question.choices.all())
            for j, choice in enumerate(choices):
                if choice_format == 'eng':
                    choice_char = chr(ord('A') + j)
                else:
                    choice_char = THAI_CHOICE_CHARS[j] if j < len(THAI_CHOICE_CHARS) else '?'
                
                choice_text = f"{choice_char}. {choice.choice_text}"
                para_choice = Paragraph(choice_text, styles['ThaiChoice'])
                w, h = para_choice.wrapOn(p, width - 2.5*inch, height)
                if y_position - h < inch:
                    p.showPage()
                    y_position = height - inch
                para_choice.drawOn(p, inch, y_position - h)
                y_position -= (h + 5)
        
        y_position -= 15
    
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

# ==============================================================================
# Word (.docx) Generation Utility
# ==============================================================================

def generate_word_exam(exam, choice_format='thai'):
    """
    Generates a Word (.docx) file for a given Exam object, including images.
    """
    document = Document()
    document.add_heading(f"ชุดข้อสอบ: {exam.exam_name}", level=1)
    document.add_paragraph(f"รายวิชา: {exam.course}")
    document.add_paragraph()
    
    for i, question in enumerate(exam.questions.all(), 1):
        p_question = document.add_paragraph(style='List Number')
        p_question.add_run(question.question_text).bold = False

        # Add image if it exists
        if question.image:
            try:
                # Add picture with a specified width (height will be scaled automatically)
                document.add_picture(question.image.path, width=Inches(4.0))
            except Exception as e:
                print(f"Error adding image to Word: {e}")
                document.add_paragraph(f"[ไม่สามารถแทรกรูปภาพ: {question.image.name}]")

        if question.question_type == 'MCQ':
            choices = list(question.choices.all())
            for j, choice in enumerate(choices):
                if choice_format == 'eng':
                    choice_char = chr(ord('A') + j)
                else:
                    choice_char = THAI_CHOICE_CHARS[j] if j < len(THAI_CHOICE_CHARS) else '?'
                
                p_choice = document.add_paragraph(f"{choice_char}. {choice.choice_text}")
                p_choice.paragraph_format.left_indent = Inches(0.5)
        
        # Add a small space after each question block
        document.add_paragraph().add_run().font.size = Pt(6)
        
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer